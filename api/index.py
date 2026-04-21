import os
import json
import re
import io
import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from google import genai
from dotenv import load_dotenv

# 1. Load Environment Variables & API Client
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("API Key Gemini tidak ditemukan di file .env!")

# Inisialisasi Client Gemini versi terbaru (SDK 2026)
client = genai.Client(api_key=GEMINI_API_KEY)

app = FastAPI(title="Word to Form API - AI Parser")
router = APIRouter(prefix="/api")

# 2. Setup CORS agar bisa diakses dari Next.js (localhost:3000)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Untuk development bisa pakai "*", untuk production ganti ke domain frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_docx(file_path):
    """Fungsi untuk mengambil teks dari paragraf dan tabel di file .docx"""
    doc = Document(file_path)
    full_text = []
    
    # Ambil teks dari paragraf biasa
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Ambil teks dari dalam tabel (penting karena soal sering dalam tabel)
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                full_text.append(" | ".join(row_data))
                
    return "\n".join(full_text)

@router.post("/convert")
async def convert_word_to_json(file: UploadFile = File(...)):
    try:
        # 1. Baca file langsung ke memory (tanpa simpan ke disk, serverless-friendly)
        content = await file.read()
        file_stream = io.BytesIO(content)

        # 2. Ekstrak teks mentah dari file Word memori
        raw_text = extract_text_from_docx(file_stream)
        
        if not raw_text.strip():
            raise HTTPException(status_code=400, detail="File Word kosong atau tidak terbaca.")

        # 3. Prompt Antigravity: Perintah sakti untuk AI
        system_instruction = """
# ROLE
Kamu adalah Expert Educational Data Architect dan Senior Parser yang spesialis dalam digitalisasi dokumen akademik Indonesia. Tugas utama kamu adalah mengekstraksi teks mentah dari Microsoft Word menjadi struktur data JSON yang valid, bersih, dan siap pakai untuk sistem Learning Management System (LMS) atau Google Forms.

# CONTEXT & INPUT
Input yang kamu terima adalah hasil ekstraksi teks dari dokumen .docx yang bisa berupa:
1. Lembar Soal Ujian (Pilihan Ganda, Checkbox, Esai).
2. Dokumen Kurikulum (Alur Tujuan Pembelajaran/ATP, Capaian Pembelajaran/CP).
3. Modul Ajar atau Silabus.

# CORE EXTRACTION RULES
1. IDENTIFIKASI METADATA: 
   Ekstrak informasi umum seperti Judul Dokumen, Mata Pelajaran, Nama Guru/Penyusun, NIP, Fase/Kelas, dan Nama Sekolah (Contoh: SMP Negeri 3 Sumenep).

2. LOGIKA PENOMORAN:
   - Kenali penomoran soal standar (1, 2, 3).
   - Kenali penomoran hierarki kurikulum (Contoh: 8.1.1, 8.2.1) dan petakan sebagai ID unik untuk butir tersebut.

3. DETEKSI TIPE PERTANYAAN/ITEM:
   - "multiple_choice": Jika terdapat satu set pilihan (A-E) dengan instruksi memilih satu jawaban.
   - "checkbox": Jika terdapat instruksi "Pilih lebih dari satu", "Pilih 2 jawaban", atau pilihan ganda kompleks.
   - "essay": Jika berupa pertanyaan terbuka atau instruksi tanpa pilihan jawaban.
   - "curriculum_objective": Khusus untuk dokumen ATP, petakan Materi dan Tujuan Pembelajaran ke dalam kategori ini.

4. HANDLING MATEMATIKA & SAINS:
   Ubah semua rumus matematika, kimia, atau fisika menjadi format LaTeX menggunakan pembungkus $. 
   Contoh: $x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$

5. DETEKSI GAMBAR & VISUAL:
   Jika dalam teks terdapat kalimat rujukan gambar (Contoh: "Perhatikan gambar berikut", "Lihat tabel di atas", "Berdasarkan grafik..."), maka:
   - Set "has_image": true.
   - Ekstrak kalimat rujukan tersebut ke dalam "image_context_hint".

6. CLEANING:
   Hapus sisa-sisa karakter sampah hasil ekstraksi tabel atau spasi berlebih. Pastikan nama orang dan gelar (S.Pd., M.Pd.) tertulis dengan benar sesuai dokumen.

# JSON OUTPUT SCHEMA
Wajib mengembalikan output dalam format JSON murni tanpa kata-kata pembuka atau penutup. 

{
  "metadata": {
    "document_title": "String",
    "subject": "String",
    "author": "String",
    "nip": "String",
    "school": "String",
    "phase_class": "String"
  },
  "content_type": "exam_paper" | "curriculum_doc",
  "items": [
    {
      "id": "String/Integer",
      "type": "multiple_choice" | "checkbox" | "essay" | "curriculum_objective",
      "category": "String (Materi/Topik)",
      "text": "String (Pertanyaan atau Isi Tujuan Pembelajaran)",
      "options": {
        "a": "String",
        "b": "String",
        "c": "String",
        "d": "String",
        "e": "String"
      } | null,
      "answer_key": "String | null",
      "has_image": Boolean,
      "image_context_hint": "String | null",
      "metadata_extra": {
        "profil_pancasila": "String | null",
        "jp_allocation": "String | null"
      }
    }
  ]
}

# ERROR HANDLING
Jika teks sangat berantakan, gunakan logika inferensi terbaikmu untuk menentukan di mana sebuah pertanyaan dimulai dan berakhir berdasarkan pola penomoran dokumen Indonesia.
"""

        # 4. Panggil Gemini AI
        model_name = "gemini-3-flash-preview" 
        
        try:
            print(f"Menggunakan model: {model_name}")
            response = client.models.generate_content(
                model=model_name,
                contents=f"{system_instruction}\n\nTEKS SUMBER DARI WORD:\n{raw_text}"
            )
            
            # Jika response berhasil
            raw_ai_response = response.text
            print("Berhasil mendapatkan respon dari AI!")
        except Exception as ai_err:
            # Jika masih error (misal kuota habis), kita cetak error spesifiknya
            print(f"AI Error: {ai_err}")
            raise HTTPException(status_code=500, detail="AI sedang sibuk atau kuota habis. Coba lagi 1 menit lagi.")

        # Tambahkan ini sebentar buat cek
        # print("Cek model yang tersedia...")
        # for m in client.models.list():
        #     print(f"- {m.name}")
        
        # 5. Parsing & Cleaning Response AI
        raw_ai_response = response.text
        
        # Ekstraksi JSON menggunakan Regex agar lebih tangguh
        clean_json_str = ""
        json_match = re.search(r'\{.*\}', raw_ai_response, re.DOTALL)
        if json_match:
            clean_json_str = json_match.group(0)
        else:
            clean_json_str = raw_ai_response.strip()
        
        try:
            parsed_json = json.loads(clean_json_str)
        except json.JSONDecodeError:
            print(f"Gagal parsing JSON. Konten mentah: {raw_ai_response}")
            raise HTTPException(status_code=500, detail="AI tidak memberikan format JSON yang valid.")

        return {
            "status": "success",
            "filename": file.filename,
            "data": parsed_json
        }

    except Exception as e:
        print(f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


app.include_router(router)

if __name__ == "__main__":
    # Jalankan server di port 8000
    uvicorn.run(app, host="0.0.0.0", port=8000)